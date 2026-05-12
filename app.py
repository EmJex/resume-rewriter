"""
智能简历改写助手 v3 - 体检 + JD匹配优化
"""

import streamlit as st
from openai import OpenAI
import io
import os
import re
import difflib
from docx import Document
from docx.shared import Pt, RGBColor

st.set_page_config(page_title="智能简历改写助手", page_icon="📝", layout="wide")

DEFAULT_BASE_URL = "https://api.openai.com/v1"
DEFAULT_MODEL = "gpt-4o"

CHECKUP_PROMPT = """请你扮演一名资深招聘官，根据我提供的简历内容，进行一次全面的专业评估。请指出简历在结构、内容、语言表达、专业性、逻辑清晰度和吸引力等方面存在的优点和不足。同时，请给出具体的改进建议，说明为什么要这样修改。

请按以下格式输出：

## 📊 简历总体评分
给出1-10分的评分和一句话总评

## ✅ 优点
- 列出简历做得好的地方

## ⚠️ 需要改进的地方
- 列出具体问题和改进建议，说明原因

## 🎯 修改优先级
按紧急程度排序，告诉用户先改什么

请用中文回答，具体、有针对性，不要泛泛而谈。"""

JD_MATCH_PROMPT = """请你作为一位熟悉目标行业的职业顾问，根据我提供的简历和目标岗位JD，深度优化我的简历。

请按以下步骤分析并输出：

## 1️⃣ JD核心要求拆解
分析JD中的核心职责、关键技能和经验要求

## 2️⃣ 匹配度分析
### 🟢 高度匹配的部分
指出简历中与JD高度匹配、但可以更突出强调的部分

### 🟡 需要补充的部分
指出简历中与JD相关但描述不够充分，需要补充或修改的部分

### 🔴 缺失的部分
指出JD要求但简历中完全没有体现的内容，给出补充建议

## 3️⃣ 工作经历和项目经验修改建议
为每段工作经历和项目提供具体修改建议，使语言和侧重点更贴合JD

## 4️⃣ ATS关键词优化
建议在技能部分增加或调整哪些关键词，以提高ATS匹配度。列出建议添加的关键词列表。

请用中文回答。"""

REWRITE_PROMPT = """请根据目标岗位JD，改写这份简历。

改写原则：
1. 从JD提取核心关键词，自然融入简历
2. 用STAR法则重写工作经验，突出与目标岗位相关的成就
3. 尽可能量化成果（提升XX%、节省XX万等）
4. 不捏造经历，只调整表达和侧重点
5. 教育背景、个人信息、时间线等事实不变
6. 语言与原始简历一致
7. 保持原始结构格式，用##标记标题，用-标记列表

只输出改写后的完整简历，不要加任何解释。"""

TIPS_PROMPT = """基于这份简历和目标岗位，给出具体的投递优化建议：

1. 📝 **求职信要点**：针对这个岗位，求职信应该重点突出什么
2. 📎 **加分材料**：作品集、推荐信、证书等建议
3. 🎤 **面试准备**：可能被问到的问题和准备方向
4. 📮 **投递策略**：投递时间、是否需要内推、如何跟进
5. 📚 **技能差距弥补**：如果有明显差距，短期内如何弥补
6. 🔑 **ATS通过技巧**：针对这个岗位的ATS优化提醒

每条建议要具体可执行，不要泛泛而谈。用中文回答。"""


def extract_text_from_pdf(file) -> str:
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file)
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    except Exception as e:
        st.error(f"PDF解析失败: {e}")
        return ""


def extract_text_from_docx(file) -> str:
    try:
        doc = Document(file)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()
    except Exception as e:
        st.error(f"Word文档解析失败: {e}")
        return ""


def create_docx(text: str) -> bytes:
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith(('- ', '• ', '* ')):
            p = doc.add_paragraph(style='List Bullet')
            _add_fmt(p, line[2:].strip())
        elif re.match(r'^\d+[\.\)]\s', line):
            p = doc.add_paragraph(style='List Number')
            _add_fmt(p, re.sub(r'^\d+[\.\)]\s', '', line).strip())
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('─' * 50).font.color.rgb = RGBColor(200, 200, 200)
        else:
            p = doc.add_paragraph()
            _add_fmt(p, line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _add_fmt(p, text):
    for part in re.split(r'(\*\*.*?\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)


def diff_html(orig, new):
    ol = orig.splitlines()
    nl = new.splitlines()
    sm = difflib.SequenceMatcher(None, ol, nl)
    h = ['<div style="font-family:sans-serif;line-height:1.8;font-size:14px;">']
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'equal':
            for l in ol[i1:i2]:
                if l.strip():
                    h.append(f'<div style="padding:2px 8px;color:#333;">{_esc(l)}</div>')
        elif tag == 'replace':
            for l in ol[i1:i2]:
                if l.strip():
                    h.append(f'<div style="padding:2px 8px;background:#fdd;text-decoration:line-through;color:#999;">🔴 {_esc(l)}</div>')
            for l in nl[j1:j2]:
                if l.strip():
                    h.append(f'<div style="padding:2px 8px;background:#dfd;color:#333;">🟢 {_esc(l)}</div>')
        elif tag == 'delete':
            for l in ol[i1:i2]:
                if l.strip():
                    h.append(f'<div style="padding:2px 8px;background:#fdd;text-decoration:line-through;color:#999;">🔴 {_esc(l)}</div>')
        elif tag == 'insert':
            for l in nl[j1:j2]:
                if l.strip():
                    h.append(f'<div style="padding:2px 8px;background:#dfd;color:#333;">🟢 {_esc(l)}</div>')
    h.append('</div>')
    return '\n'.join(h)


def _esc(t):
    return t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


def change_stats(orig, new):
    ol = [l for l in orig.splitlines() if l.strip()]
    nl = [l for l in new.splitlines() if l.strip()]
    sm = difflib.SequenceMatcher(None, ol, nl)
    eq = mod = add = rm = 0
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'equal': eq += i2-i1
        elif tag == 'replace': mod += max(i2-i1, j2-j1)
        elif tag == 'delete': rm += i2-i1
        elif tag == 'insert': add += j2-j1
    total = eq+mod+add+rm
    rate = ((mod+add+rm)/total*100) if total else 0
    return {"rate": rate, "mod": mod, "add": add, "rm": rm}


def call_llm(api_key, base_url, model, system, user_msg, max_tokens=4096):
    client = OpenAI(api_key=api_key, base_url=base_url)
    r = client.chat.completions.create(
        model=model, max_tokens=max_tokens,
        messages=[{"role": "system", "content": system}, {"role": "user", "content": user_msg}]
    )
    return r.choices[0].message.content


st.title("📝 智能简历改写助手")
st.caption("第一步体检诊断 → 第二步针对JD一键优化 → 下载投递")

with st.sidebar:
    st.header("⚙️ 设置")
    api_key = st.text_input("API Key", type="password", value=os.environ.get("OPENAI_API_KEY", ""))
    base_url = st.text_input("API Base URL", value=os.environ.get("OPENAI_BASE_URL", DEFAULT_BASE_URL))
    model = st.text_input("模型名称", value=os.environ.get("MODEL_NAME", DEFAULT_MODEL))

st.subheader("📄 上传简历")
upload_method = st.radio("输入方式", ["上传文件", "粘贴文本"], horizontal=True, label_visibility="collapsed")
resume_text = ""
if upload_method == "上传文件":
    f = st.file_uploader("上传简历", type=["pdf", "docx", "doc", "txt"])
    if f:
        if f.name.endswith('.pdf'): resume_text = extract_text_from_pdf(f)
        elif f.name.endswith(('.docx','.doc')): resume_text = extract_text_from_docx(f)
        elif f.name.endswith('.txt'): resume_text = f.read().decode('utf-8')
        if resume_text:
            with st.expander("📋 简历内容预览", expanded=False):
                st.text(resume_text)
else:
    resume_text = st.text_area("粘贴简历内容", height=300, placeholder="在此粘贴你的简历全文...")

st.divider()

st.subheader("🩺 第一步：简历体检")
st.caption("先让AI诊断简历问题，看看哪里需要改进")

checkup_btn = st.button("🔍 开始体检", disabled=not (resume_text and api_key), use_container_width=True)

if checkup_btn and resume_text and api_key:
    with st.spinner("正在体检..."):
        try:
            checkup = call_llm(api_key, base_url, model, CHECKUP_PROMPT, f"我的简历内容如下：\n{resume_text}")
            st.session_state['checkup'] = checkup
        except Exception as e:
            st.error(f"体检失败: {e}")

if 'checkup' in st.session_state:
    st.markdown(st.session_state['checkup'])

st.divider()

st.subheader("🎯 第二步：针对目标JD一键优化")
st.caption("粘贴目标岗位描述，AI帮你定制简历")

jd_text = st.text_area("粘贴目标岗位描述（JD）", height=250, placeholder="在此粘贴目标岗位的招聘描述...")

col_a, col_b = st.columns(2)
with col_a:
    optimize_btn = st.button("🚀 分析 + 改写", type="primary", disabled=not (resume_text and jd_text and api_key), use_container_width=True)
with col_b:
    tips_btn = st.button("💡 投递优化建议", disabled=not (resume_text and jd_text and api_key), use_container_width=True)

if optimize_btn and resume_text and jd_text and api_key:
    input_text = f"## 我的简历\n{resume_text}\n\n## 目标岗位JD\n{jd_text}"

    with st.status("🔍 分析JD匹配度...", expanded=True) as status:
        try:
            match_analysis = call_llm(api_key, base_url, model, JD_MATCH_PROMPT, input_text, max_tokens=3000)
            st.session_state['match_analysis'] = match_analysis
            st.markdown(match_analysis)
            status.update(label="✅ 分析完成", state="complete")
        except Exception as e:
            st.error(f"分析失败: {e}")
            st.stop()

    with st.status("✍️ 正在改写简历...", expanded=False) as status:
        try:
            rewritten = call_llm(api_key, base_url, model, REWRITE_PROMPT, input_text + "\n\n请改写简历。")
            st.session_state['rewritten'] = rewritten
            status.update(label="✅ 改写完成", state="complete")
        except Exception as e:
            st.error(f"改写失败: {e}")
            st.stop()

if 'rewritten' in st.session_state and resume_text:
    rewritten = st.session_state['rewritten']

    st.divider()

    stats = change_stats(resume_text, rewritten)
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("📊 修改幅度", f"{stats['rate']:.0f}%")
    c2.metric("✏️ 修改", stats['mod'])
    c3.metric("➕ 新增", stats['add'])
    c4.metric("➖ 删除", stats['rm'])

    tab1, tab2 = st.tabs(["🔍 修改对比", "📄 改写全文"])
    with tab1:
        st.markdown("**🔴 删除线 = 原文  |  🟢 绿色 = 改写后**")
        st.html(diff_html(resume_text, rewritten))
    with tab2:
        st.markdown(rewritten)

    st.divider()
    d1, d2 = st.columns(2)
    with d1:
        st.download_button("📥 下载 Word", create_docx(rewritten), "改写简历.docx",
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
    with d2:
        st.download_button("📥 下载纯文本", rewritten.encode('utf-8'), "改写简历.txt", "text/plain", use_container_width=True)

if tips_btn and resume_text and jd_text and api_key:
    with st.spinner("生成投递优化建议..."):
        try:
            tips = call_llm(api_key, base_url, model, TIPS_PROMPT,
                          f"## 简历\n{resume_text}\n\n## 目标岗位\n{jd_text}")
            st.session_state['tips'] = tips
        except Exception as e:
            st.error(f"生成失败: {e}")

if 'tips' in st.session_state:
    st.divider()
    st.subheader("💡 投递优化建议")
    st.markdown(st.session_state['tips'])
